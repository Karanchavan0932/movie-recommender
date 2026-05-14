from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_color(doc, text, level, color_rgb=(0, 102, 204)):
    """Add a heading with custom color"""
    heading = doc.add_heading(text, level=level)
    heading.style = f'Heading {level}'
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
    return heading

def add_colored_paragraph(doc, text, color_rgb=(0, 0, 0), bold=False, size=11):
    """Add a paragraph with custom color and formatting"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
        run.font.bold = bold
        run.font.size = Pt(size)
    return p

def shade_paragraph(paragraph, color):
    """Add background color to paragraph"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    paragraph._element.get_or_add_pPr().append(shading_elm)

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('🎬 Movie Recommender System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(220, 20, 60)
    run.font.size = Pt(28)

# Subtitle
subtitle = doc.add_paragraph('Project Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# 1. INTRODUCTION
add_heading_with_color(doc, '1. Introduction', 1, (0, 102, 204))
intro_text = """The Movie Recommender System is an intelligent application designed to suggest movies based on user preferences. By leveraging advanced machine learning algorithms and content-based filtering techniques, this system analyzes movie genres and provides personalized recommendations to enhance the user's movie-watching experience.

The primary objective of this project is to create a user-friendly interface that enables movie enthusiasts to discover new films similar to their favorites. The system processes movie metadata and computes similarity scores between movies to generate accurate recommendations in real-time.

This project demonstrates practical applications of data science, machine learning, and modern web technologies in building a scalable recommendation engine suitable for deployment in production environments."""

doc.add_paragraph(intro_text)

doc.add_paragraph()

# 2. WORKING / HOW IT WORKS
add_heading_with_color(doc, '2. Working / System Architecture', 1, (0, 102, 204))

doc.add_paragraph('The Movie Recommender System operates through the following process:', style='List Bullet')

working_steps = [
    'Data Loading: The system loads movie data from a CSV file containing movie titles and genres.',
    'Data Preprocessing: Missing values are handled gracefully, and data is cleaned to ensure consistency.',
    'Feature Vectorization: Movie genres are converted into numerical vectors using CountVectorizer from scikit-learn.',
    'Similarity Calculation: Cosine similarity is computed between all movies based on their genre vectors.',
    'Recommendation Generation: When a user enters a movie name, the system retrieves the top 5 most similar movies.',
    'Result Display: Recommendations are displayed in an intuitive, user-friendly interface built with Streamlit.'
]

for step in working_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph()

# System Flow Diagram Description
add_heading_with_color(doc, '2.1 System Flow', 2, (0, 102, 204))
flow_description = """
1. User Input → 2. Search Movie Database → 3. Extract Movie Features (Genre)
   ↓
4. Compute Similarity Scores → 5. Rank Movies by Similarity → 6. Return Top 5 Recommendations
   ↓
7. Display Results to User
"""
doc.add_paragraph(flow_description)

doc.add_paragraph()

# 3. TECHNOLOGIES USED
add_heading_with_color(doc, '3. Technologies Used', 1, (0, 102, 204))

technologies = [
    ('Python 3.x', 'Core programming language for backend development'),
    ('Streamlit', 'Web framework for building interactive data applications with minimal code'),
    ('Pandas', 'Data manipulation and analysis library for handling CSV files and dataframes'),
    ('NumPy', 'Numerical computing library for mathematical operations and array handling'),
    ('scikit-learn', 'Machine learning library providing CountVectorizer and cosine similarity functions'),
    ('CSV Format', 'Data storage format for movie information'),
]

# Create table for technologies
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Technology'
header_cells[1].text = 'Purpose'

# Format header
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

for tech, purpose in technologies:
    row_cells = table.add_row().cells
    row_cells[0].text = tech
    row_cells[1].text = purpose

doc.add_paragraph()

# 4. ALGORITHM USED
add_heading_with_color(doc, '4. Algorithm: Cosine Similarity & Content-Based Filtering', 1, (0, 102, 204))

add_heading_with_color(doc, '4.1 Overview', 2, (0, 102, 204))
algo_overview = """The Movie Recommender System uses Content-Based Filtering with Cosine Similarity as its core algorithm. This approach analyzes movie features (specifically genres) and recommends movies that are most similar to the user's input movie."""

doc.add_paragraph(algo_overview)

add_heading_with_color(doc, '4.2 Algorithm Steps', 2, (0, 102, 204))

algo_steps = [
    'Feature Extraction: Extract genre information from each movie in the dataset.',
    'Vectorization: Convert genre text into numerical vectors using CountVectorizer (Bag-of-Words model).',
    'Similarity Computation: Calculate cosine similarity between the input movie and all other movies.',
    'Ranking: Sort movies by similarity scores in descending order.',
    'Selection: Return the top N (default 5) most similar movies excluding the input movie itself.'
]

for step in algo_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph()

add_heading_with_color(doc, '4.3 Cosine Similarity Formula', 2, (0, 102, 204))
formula_para = doc.add_paragraph()
formula_para.add_run('Cosine Similarity (A, B) = ').font.size = Pt(12)
formula_para.add_run('(A · B) / (||A|| × ||B||)').font.size = Pt(12)
formula_para.add_run('  , where value ranges from 0 to 1').font.size = Pt(11)
formula_para.add_run('\n• A · B = dot product of vectors\n• ||A|| and ||B|| = magnitude of vectors\n• Value closer to 1 = higher similarity').font.size = Pt(11)

doc.add_paragraph()

add_heading_with_color(doc, '4.4 Why Cosine Similarity?', 2, (0, 102, 204))
reasons = [
    'Efficient: Computationally fast and scalable for large datasets',
    'Effective: Works well with high-dimensional sparse data (genre vectors)',
    'Normalized: Similarity scores are bounded between 0 and 1, making interpretation straightforward',
    'Simple: Easy to implement and understand',
    'Content-Based: Does not require user ratings or collaborative data'
]

for reason in reasons:
    doc.add_paragraph(reason, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, '4.5 Example', 2, (0, 102, 204))
example_text = """
Movie 1: "The Matrix" → Genres: Action, Sci-Fi
Movie 2: "Inception" → Genres: Action, Sci-Fi, Thriller
Movie 3: "Titanic" → Genres: Romance, Drama

When user searches "The Matrix":
- Similarity with Inception: 0.89 (High - shares Action, Sci-Fi genres)
- Similarity with Titanic: 0.25 (Low - no common genres)
- Result: "Inception" will be recommended as the top suggestion
"""
doc.add_paragraph(example_text)

doc.add_paragraph()

# 5. FEATURES & CAPABILITIES
add_heading_with_color(doc, '5. Key Features', 1, (0, 102, 204))

features = [
    'Real-time Recommendations: Instant movie suggestions based on genre similarity',
    'Error Handling: Graceful handling of missing files and invalid inputs',
    'Data Caching: Optimized performance through intelligent caching mechanisms',
    'User-Friendly Interface: Intuitive Streamlit-based UI with emojis and clear guidance',
    'Scalability: Architecture supports easy scaling and deployment to cloud platforms',
    'Case-Insensitive Search: Flexible movie name input without exact case matching'
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph()

# 6. USAGE INSTRUCTIONS
add_heading_with_color(doc, '6. How to Use', 1, (0, 102, 204))

add_heading_with_color(doc, '6.1 Installation', 2, (0, 102, 204))
install_steps = [
    'Create a virtual environment: python -m venv venv',
    'Activate virtual environment: venv\\Scripts\\activate (Windows)',
    'Install dependencies: pip install -r requirements.txt',
    'Run the application: streamlit run app.py'
]

for step in install_steps:
    doc.add_paragraph(step, style='List Number')

add_heading_with_color(doc, '6.2 Running the Application', 2, (0, 102, 204))
run_text = """1. Execute the command: streamlit run app.py
2. The application will open in your default browser at http://localhost:8501
3. Enter a movie name in the text field
4. Click the "Recommend" button
5. View the 5 most similar movies"""

doc.add_paragraph(run_text)

doc.add_paragraph()

# 7. ADVANTAGES & APPLICATIONS
add_heading_with_color(doc, '7. Advantages & Potential Applications', 1, (0, 102, 204))

add_heading_with_color(doc, '7.1 Advantages', 2, (0, 102, 204))
advantages = [
    'Low Latency: Fast recommendations without requiring complex database queries',
    'No Cold Start Problem for Items: Works well even for new movies with genres',
    'Transparency: Users can easily understand why a movie is recommended',
    'No User Data Required: Does not require user ratings or personal information',
    'Easy Maintenance: Simple to update with new movie data'
]

for adv in advantages:
    doc.add_paragraph(adv, style='List Bullet')

add_heading_with_color(doc, '7.2 Potential Applications', 2, (0, 102, 204))
applications = [
    'Entertainment Platforms: Netflix, Amazon Prime, Disney+ style recommendation systems',
    'Movie Theaters: Interactive kiosks suggesting movies to customers',
    'Educational Platforms: Recommending educational documentaries or films',
    'Content Curation: Automated movie list generation for streaming services',
    'Social Networks: Friend-based movie recommendations'
]

for app in applications:
    doc.add_paragraph(app, style='List Bullet')

doc.add_paragraph()

# 8. LIMITATIONS
add_heading_with_color(doc, '8. Limitations & Future Improvements', 1, (0, 102, 204))

add_heading_with_color(doc, '8.1 Current Limitations', 2, (0, 102, 204))
limitations = [
    'Genre-Only Analysis: Recommendations based solely on genres, ignoring other factors like cast, director, year',
    'Limited Features: No consideration of user ratings or popularity scores',
    'Scalability: May become slower with very large datasets (millions of movies)',
    'New Movie Problem: Requires manual genre entry for new movies',
    'No User Personalization: Cannot adapt recommendations to individual user preferences over time'
]

for lim in limitations:
    doc.add_paragraph(lim, style='List Bullet')

add_heading_with_color(doc, '8.2 Future Enhancements', 2, (0, 102, 204))
enhancements = [
    'Hybrid Approach: Combine content-based filtering with collaborative filtering',
    'Additional Features: Include director, cast, release year, and user ratings',
    'Deep Learning: Implement neural networks for better feature representation',
    'User Profiles: Add user registration and personalized recommendation history',
    'Ratings & Feedback: Allow users to rate recommendations for continuous improvement',
    'Database Integration: Use SQL/NoSQL databases for better scalability'
]

for enh in enhancements:
    doc.add_paragraph(enh, style='List Bullet')

doc.add_paragraph()

# 9. CONCLUSION
add_heading_with_color(doc, '9. Conclusion', 1, (0, 102, 204))

conclusion_text = """The Movie Recommender System demonstrates a practical implementation of machine learning concepts in a real-world application. By leveraging content-based filtering and cosine similarity algorithms, the system successfully recommends movies that align with user preferences based on genre similarity.

Key Achievements:
• Built an efficient, user-friendly recommendation system using Python and Streamlit
• Successfully implemented the cosine similarity algorithm for accurate recommendations
• Created a scalable architecture ready for deployment and enhancement
• Demonstrated effective error handling and data preprocessing techniques

This project serves as an excellent foundation for more advanced recommendation systems. By integrating additional features, collaborative filtering, and machine learning models, the system can be further enhanced to provide more personalized and accurate recommendations.

The Movie Recommender System represents a bridge between theoretical machine learning concepts and practical, production-ready applications. It showcases how simple yet powerful algorithms can create valuable solutions that enhance user experiences in the entertainment industry.

Potential Next Steps:
1. Integrate user ratings and feedback mechanisms
2. Expand features to include cast, director, and release year
3. Implement collaborative filtering for enhanced personalization
4. Deploy on cloud platforms (Streamlit Cloud, AWS, Azure)
5. Add database backend for better scalability and data management

This project proves that with the right combination of technologies and algorithms, creating intelligent recommendation systems is not only feasible but also practical and effective."""

doc.add_paragraph(conclusion_text)

doc.add_paragraph()
doc.add_paragraph()

# Footer / Project Information
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('Movie Recommender System - Complete Project Documentation')
footer_run.font.size = Pt(10)
footer_run.font.italic = True
footer_run.font.color.rgb = RGBColor(128, 128, 128)

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer2_run = footer2.add_run('Built with Python, Streamlit, Pandas, NumPy, and scikit-learn')
footer2_run.font.size = Pt(9)
footer2_run.font.color.rgb = RGBColor(128, 128, 128)

# Save the document
output_path = 'Movie_Recommender_System_Documentation.docx'
doc.save(output_path)
print(f"✅ Document created successfully: {output_path}")
